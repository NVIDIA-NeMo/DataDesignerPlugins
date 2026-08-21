# SPDX-FileCopyrightText: Copyright (c) 2026 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

import json
from pathlib import Path

import pandas as pd
import pytest
from data_designer.config import ExpressionColumnConfig
from data_designer.config.config_builder import DataDesignerConfigBuilder
from data_designer.config.seed_source_dataframe import DataFrameSeedSource
from data_designer.engine.testing.utils import assert_valid_plugin
from data_designer.interface.data_designer import DataDesigner
from docx import Document
from pydantic import ValidationError

from data_designer_docx.config import DocxProcessorConfig
from data_designer_docx.impl import DocxProcessor
from data_designer_docx.plugin import plugin
from data_designer_docx.render import normalize_rows, render_document, safe_filename
from data_designer_docx.schema import DocSection, DocTable, WordDocument


def test_valid_plugin() -> None:
    assert_valid_plugin(plugin)


def make_document(title: str = "Access Control Standard") -> WordDocument:
    """Build a small but complete document for rendering tests."""
    return WordDocument(
        title=title,
        subtitle="Information Security",
        summary="This standard defines how access is granted, reviewed, and revoked.",
        sections=[
            DocSection(heading="Scope", paragraphs=["Applies to all systems."], bullets=["Production"]),
            DocSection(heading="Roles", paragraphs=["The owner attests quarterly."]),
        ],
        key_data=DocTable(
            caption="Review Cadence",
            columns=["Role", "Action", "Frequency"],
            rows=[["Owner", "Attest", "Quarterly"]],
        ),
    )


class BoundDocxProcessor(DocxProcessor):
    """A processor bound to an explicit dataset path.

    Lets the collision and path-containment logic be tested directly, without
    standing up a ResourceProvider.
    """

    _base_dataset_path: Path

    @property
    def base_dataset_path(self) -> Path:
        return self._base_dataset_path


def build_processor(tmp_path: Path, **overrides: object) -> BoundDocxProcessor:
    """Construct a BoundDocxProcessor over a temporary dataset directory."""
    dataset_path = tmp_path / "dataset"
    dataset_path.mkdir(exist_ok=True)

    processor = BoundDocxProcessor.__new__(BoundDocxProcessor)
    processor._base_dataset_path = dataset_path
    processor._config = DocxProcessorConfig(name="docs", document_column="document", **overrides)
    processor._initialize()
    return processor


class TestDocxProcessorConfig:
    def test_defaults(self) -> None:
        config = DocxProcessorConfig(name="docs", document_column="document")
        assert config.processor_type == "docx"
        assert config.output_subdir == "documents"
        assert config.output_path_column == "docx_path"

    def test_default_filename_needs_no_dataset_column(self) -> None:
        """A minimal config must be renderable without inventing an unrelated id column."""
        config = DocxProcessorConfig(name="docs", document_column="document")
        assert "{{" not in config.filename_template

    @pytest.mark.parametrize(
        "reserved",
        [
            "processors-files",
            "parquet-files",
            "dropped-columns-parquet-files",
            # Deleted by Data Designer on resume, which would orphan every docx_path.
            "tmp-partial-parquet-files",
            "images",
            "./processors-files",
            "foo/../parquet-files",
            "nested/images",
        ],
    )
    def test_reserved_output_subdir_is_rejected(self, reserved: str) -> None:
        """Data Designer reads or deletes these folders, so documents must not go there."""
        with pytest.raises(ValidationError):
            DocxProcessorConfig(name="docs", document_column="document", output_subdir=reserved)

    @pytest.mark.parametrize("escaping", ["../outside", "/private/tmp/outside", "a/../../outside", ""])
    def test_escaping_output_subdir_is_rejected(self, escaping: str) -> None:
        with pytest.raises(ValidationError):
            DocxProcessorConfig(name="docs", document_column="document", output_subdir=escaping)

    @pytest.mark.parametrize("bad_name", ["../../outside", "/abs", "nested/name", "..", "processors-files"])
    def test_unsafe_processor_name_is_rejected(self, bad_name: str) -> None:
        """The processor name becomes a directory, so it is a traversal route too."""
        with pytest.raises(ValidationError):
            DocxProcessorConfig(name=bad_name, document_column="document")

    def test_nested_output_subdir_is_allowed(self) -> None:
        config = DocxProcessorConfig(name="docs", document_column="document", output_subdir="out/word")
        assert config.output_subdir == "out/word"


class TestSafeFilename:
    @pytest.mark.parametrize(
        ("raw", "expected"),
        [
            ("POL-1 Access Control", "POL-1-Access-Control.docx"),
            ("already.docx", "already.docx"),
            ("../../etc/passwd", "etc-passwd.docx"),
            ("///", "document.docx"),
        ],
    )
    def test_sanitizes(self, raw: str, expected: str) -> None:
        assert safe_filename(raw) == expected


class TestNormalizeRows:
    def test_pads_short_rows(self) -> None:
        """Structured outputs constrain JSON shape, not cell counts, so rows arrive ragged."""
        table = DocTable(caption="c", columns=["a", "b", "c"], rows=[["1", "2"]])
        assert normalize_rows(table) == [["1", "2", ""]]

    def test_truncates_long_rows(self) -> None:
        table = DocTable(caption="c", columns=["a", "b"], rows=[["1", "2", "3"]])
        assert normalize_rows(table) == [["1", "2"]]


class TestRenderDocument:
    def test_writes_readable_docx(self, tmp_path: Path) -> None:
        path = render_document(
            make_document(),
            tmp_path / "out.docx",
            metadata={"Document ID": "POL-1"},
            footer_text="Internal",
            core_properties={"author": "Jane Doe", "category": "Standard"},
        )

        rendered = Document(str(path))
        headings = [p.text for p in rendered.paragraphs if p.style.name.startswith(("Title", "Heading"))]
        assert headings[0] == "Access Control Standard"
        assert "1. Scope" in headings
        assert rendered.core_properties.author == "Jane Doe"
        assert rendered.sections[0].footer.paragraphs[0].text == "Internal"

    def test_metadata_and_key_data_tables(self, tmp_path: Path) -> None:
        path = render_document(make_document(), tmp_path / "out.docx", metadata={"Owner": "Jane"})

        rendered = Document(str(path))
        assert len(rendered.tables) == 2
        assert [cell.text for cell in rendered.tables[1].rows[0].cells] == ["Role", "Action", "Frequency"]

    def test_number_sections_disabled(self, tmp_path: Path) -> None:
        path = render_document(make_document(), tmp_path / "out.docx", number_sections=False)

        rendered = Document(str(path))
        headings = [p.text for p in rendered.paragraphs if p.style.name.startswith("Heading")]
        assert "Scope" in headings
        assert "1. Scope" not in headings

    def test_template_styles_are_inherited(self, tmp_path: Path) -> None:
        template_path = tmp_path / "template.docx"
        template = Document()
        template.styles["Normal"].font.name = "Georgia"
        template.save(str(template_path))

        path = render_document(make_document(), tmp_path / "out.docx", template_path=template_path)

        assert Document(str(path)).styles["Normal"].font.name == "Georgia"


class TestDocxProcessorPreviewIntegration:
    """Run the processor through Data Designer using seed data, so no model is needed."""

    @pytest.fixture()
    def seed_df(self) -> pd.DataFrame:
        return pd.DataFrame(
            {
                "doc_id": ["POL-1", "POL-2"],
                "document": [make_document("First").model_dump_json(), make_document("Second").model_dump_json()],
            }
        )

    def build(self, seed_df: pd.DataFrame, **overrides: object) -> DataDesignerConfigBuilder:
        builder = DataDesignerConfigBuilder()
        builder.with_seed_dataset(DataFrameSeedSource(df=seed_df))
        # Data Designer's profiler requires at least one generated column, and a
        # seed-plus-processor config has none.
        builder.add_column(ExpressionColumnConfig(name="doc_label", expr="{{ doc_id }}"))
        builder.add_processor(
            DocxProcessorConfig(
                name="word-documents",
                document_column="document",
                filename_template="{{ doc_id }}.docx",
                **overrides,
            )
        )
        return builder

    def test_preview_writes_documents(self, seed_df: pd.DataFrame, tmp_path: Path) -> None:
        artifact_path = tmp_path / "artifacts"
        artifact_path.mkdir()

        result = DataDesigner(artifact_path=artifact_path).preview(self.build(seed_df), num_records=2)

        assert "docx_path" in result.dataset.columns
        paths = sorted(artifact_path.rglob("*.docx"))
        assert [path.name for path in paths] == ["POL-1.docx", "POL-2.docx"]
        assert Document(str(paths[0])).paragraphs[0].text == "First"

    def test_output_path_column_resolves(self, seed_df: pd.DataFrame, tmp_path: Path) -> None:
        """The stored path is dataset-relative, keeping the dataset portable.

        Uses create() rather than preview() because only DatasetCreationResults
        exposes artifact_storage, which is how a caller resolves the path.
        """
        artifact_path = tmp_path / "artifacts"
        artifact_path.mkdir()

        results = DataDesigner(artifact_path=artifact_path).create(
            self.build(seed_df), num_records=2, dataset_name="documents-test"
        )

        relative = results.load_dataset()["docx_path"].iloc[0]
        assert relative.startswith("documents/word-documents/")
        assert (results.artifact_storage.base_dataset_path / relative).is_file()

    def test_invalid_document_is_skipped(self, tmp_path: Path) -> None:
        """A malformed row yields a null path instead of failing the whole batch."""
        artifact_path = tmp_path / "artifacts"
        artifact_path.mkdir()
        seed_df = pd.DataFrame(
            {
                "doc_id": ["POL-1", "POL-2"],
                "document": [make_document("Good").model_dump_json(), json.dumps({"nope": True})],
            }
        )

        result = DataDesigner(artifact_path=artifact_path).preview(self.build(seed_df), num_records=2)

        assert result.dataset["docx_path"].isna().sum() == 1
        assert len(list(artifact_path.rglob("*.docx"))) == 1


class TestStructuredStringPreservation:
    """Regression tests for JSON-decoding of already-decoded document mappings.

    Data Designer's recursive `deserialize_json_values` rewrites string leaves that
    look like JSON scalars, so `"30"` becomes `30` and `"true"` becomes `True`.
    Those values are exactly what key-data tables carry, and `WordDocument` rejects
    them, which silently dropped the row.
    """

    def numeric_document(self) -> dict:
        document = make_document().model_dump()
        document["key_data"]["columns"] = ["Threshold", "Enabled", "Notes"]
        document["key_data"]["rows"] = [["30", "true", "null"]]
        return document

    def test_scalar_looking_strings_survive_as_mapping(self, tmp_path: Path) -> None:
        artifact_path = tmp_path / "artifacts"
        artifact_path.mkdir()
        seed_df = pd.DataFrame({"doc_id": ["POL-1"], "document": [json.dumps(self.numeric_document())]})

        builder = DataDesignerConfigBuilder()
        builder.with_seed_dataset(DataFrameSeedSource(df=seed_df))
        builder.add_column(ExpressionColumnConfig(name="doc_label", expr="{{ doc_id }}"))
        builder.add_processor(
            DocxProcessorConfig(name="docs", document_column="document", filename_template="{{ doc_id }}.docx")
        )

        result = DataDesigner(artifact_path=artifact_path).preview(builder, num_records=1)

        assert result.dataset["docx_path"].notna().all(), "row was skipped instead of rendered"
        written = sorted(artifact_path.rglob("*.docx"))
        assert len(written) == 1
        cells = [cell.text for cell in Document(str(written[0])).tables[-1].rows[1].cells]
        assert cells == ["30", "true", "null"], f"string leaves were coerced: {cells}"


class TestFilenameCollisions:
    def make_processor_dir(self, tmp_path: Path) -> Path:
        output_dir = tmp_path / "documents" / "docs"
        output_dir.mkdir(parents=True)
        return output_dir

    def test_case_insensitive_collision(self, tmp_path: Path) -> None:
        """macOS and Windows treat A.docx and a.docx as the same file."""
        processor = build_processor(tmp_path, filename_template="{{ doc_id }}.docx")
        processor.seed_used_filenames(self.make_processor_dir(tmp_path))

        assert processor.unique_filename("Report") == "Report.docx"
        assert processor.unique_filename("report") == "report-1.docx"

    def test_seeds_from_existing_files(self, tmp_path: Path) -> None:
        """A resumed run must not overwrite documents written before the resume."""
        output_dir = self.make_processor_dir(tmp_path)
        (output_dir / "same.docx").write_bytes(b"existing")

        processor = build_processor(tmp_path, filename_template="same.docx")
        processor.seed_used_filenames(output_dir)

        assert processor.unique_filename("same") == "same-1.docx"

    def test_output_dir_stays_inside_dataset(self, tmp_path: Path) -> None:
        processor = build_processor(tmp_path)
        assert processor.output_dir.is_relative_to((tmp_path / "dataset").resolve())


class TestFooterTargeting:
    def test_footer_applies_to_every_section(self, tmp_path: Path) -> None:
        """Generated content lands in the template's last section, not the first."""
        template_path = tmp_path / "two-section-template.docx"
        template = Document()
        template.add_section()
        for index, section in enumerate(template.sections):
            section.footer.is_linked_to_previous = False
            section.footer.paragraphs[0].text = f"template-{index}"
        template.save(str(template_path))

        path = render_document(
            make_document(), tmp_path / "out.docx", template_path=template_path, footer_text="override"
        )

        footers = [section.footer.paragraphs[0].text for section in Document(str(path)).sections]
        assert set(footers) == {"override"}, footers
