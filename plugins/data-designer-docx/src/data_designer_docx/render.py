# SPDX-FileCopyrightText: Copyright (c) 2026 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0

"""Pure ``.docx`` rendering. No Data Designer imports on purpose.

Keeping the renderer free of engine imports means you can iterate on layout in a
REPL against a hand-written ``WordDocument`` and never spend a token.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from data_designer_docx.schema import DocTable, WordDocument

UNSAFE_FILENAME_CHARS = re.compile(r"[^A-Za-z0-9._-]+")

# "Table Grid" ships with the python-docx default template. Corporate templates
# usually define their own; see `table_style` on the processor config.
DEFAULT_TABLE_STYLE = "Table Grid"


def safe_filename(name: str, *, suffix: str = ".docx", max_length: int = 120) -> str:
    """Turn a rendered filename template into something safe to write to disk."""
    stem = UNSAFE_FILENAME_CHARS.sub("-", name.strip()).strip("-._")
    if stem.lower().endswith(suffix.lower()):
        stem = stem[: -len(suffix)]
    if not stem:
        stem = "document"
    return stem[:max_length] + suffix


def normalize_rows(table: DocTable) -> list[list[str]]:
    """Pad or truncate every row to the header width.

    Structured outputs constrain the *shape* of the JSON, not the arithmetic
    inside it — a model asked for four columns will occasionally hand back a row
    with three cells. Fixing it here is cheaper than a retry.
    """
    width = len(table.columns)
    normalized = []
    for row in table.rows:
        cells = [str(cell) for cell in row][:width]
        cells.extend([""] * (width - len(cells)))
        normalized.append(cells)
    return normalized


def add_metadata_table(doc: Any, metadata: dict[str, str], style: str) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = style
    for key, value in metadata.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    doc.add_paragraph()


def add_data_table(doc: Any, table: DocTable, style: str) -> None:
    doc.add_heading(table.caption, level=2)
    rendered = doc.add_table(rows=1, cols=max(len(table.columns), 1))
    rendered.style = style
    header_cells = rendered.rows[0].cells
    for idx, column in enumerate(table.columns):
        header_cells[idx].text = str(column)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in normalize_rows(table):
        cells = rendered.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def render_document(
    document: WordDocument,
    output_path: str | Path,
    *,
    metadata: dict[str, str] | None = None,
    template_path: str | Path | None = None,
    footer_text: str | None = None,
    table_style: str = DEFAULT_TABLE_STYLE,
    number_sections: bool = True,
    core_properties: dict[str, str] | None = None,
) -> Path:
    """Render a :class:`WordDocument` to a ``.docx`` file.

    Args:
        document: The structured document to render.
        output_path: Where to write the ``.docx``.
        metadata: Optional key/value pairs rendered as a front-matter table.
        template_path: Optional ``.docx`` whose styles, header, and footer are
            inherited. The template should contain styles only — any body
            content in it will appear above the generated content.
        footer_text: Optional footer applied to the first section.
        table_style: Table style name. Must exist in the (template) document.
        number_sections: Prefix section headings with ``1.``, ``2.``, ...
        core_properties: Optional Word core properties (author, category, ...).

    Returns:
        The path that was written.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path)) if template_path else Document()

    doc.add_heading(document.title, level=0)

    subtitle = doc.add_paragraph(document.subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle.runs:
        run.italic = True
        run.font.size = Pt(12)

    if metadata:
        add_metadata_table(doc, metadata, table_style)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(document.summary)

    for index, section in enumerate(document.sections, start=1):
        heading = f"{index}. {section.heading}" if number_sections else section.heading
        doc.add_heading(heading, level=1)
        for paragraph in section.paragraphs:
            doc.add_paragraph(paragraph)
        for bullet in section.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    add_data_table(doc, document.key_data, table_style)

    if footer_text:
        footer_paragraph = doc.sections[0].footer.paragraphs[0]
        footer_paragraph.text = footer_text
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Word core properties travel with the file. Downstream extraction and
    # classification pipelines read them, so it is worth filling them in.
    if core_properties:
        props = doc.core_properties
        for key, value in core_properties.items():
            if hasattr(props, key) and value is not None:
                setattr(props, key, str(value))

    doc.save(str(output_path))
    return output_path
